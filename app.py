import os
import time
import json
import hashlib
import concurrent.futures
import difflib
from datetime import datetime
from io import BytesIO

import streamlit as st
from openai import OpenAI
import requests
import pandas as pd
from docx import Document
from docx.shared import Pt

# =========================
# 📁 설정 & 디렉토리
# =========================
REPORT_DIR = "reports"
os.makedirs(REPORT_DIR, exist_ok=True)

# ─────────────────────────────────────────────
# 페이지 설정 & 테마 상태
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="AI Citation Analyzer",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="expanded",
)

if "dark_mode" not in st.session_state:
    st.session_state["dark_mode"] = False

_dark = st.session_state["dark_mode"]

# ─────────────────────────────────────────────
# 오리지널 UI 디자인 (V15 CSS 복구)
# ─────────────────────────────────────────────
if _dark:
    _bg, _bg2, _card, _border, _text, _text_muted = "#0F0F0F", "#1A1A1A", "#1E1E1E", "#333333", "#F0F0F0", "#999999"
    _header_gr = "linear-gradient(135deg,#1A1A1A 0%,#2A2A2A 60%,#3A3A3A 100%)"
    _sidebar_gr = "linear-gradient(180deg,#0F0F0F 0%,#1A1A1A 50%,#222222 100%)"
    _btn_gr = "linear-gradient(135deg,#333333,#555555)"
else:
    _bg, _bg2, _card, _border, _text, _text_muted = "#F5F5F5", "#EEEEEE", "#FFFFFF", "#DDDDDD", "#111111", "#555555"
    _header_gr = "linear-gradient(135deg,#111111 0%,#333333 60%,#555555 100%)"
    _sidebar_gr = "linear-gradient(180deg,#111111 0%,#222222 50%,#333333 100%)"
    _btn_gr = "linear-gradient(135deg,#111111,#444444)"

st.markdown(f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@300;400;500;600;700;800&display=swap');
:root {{ --bg: {_bg}; --card: {_card}; --border: {_border}; --text: {_text}; --text_muted: {_text_muted}; }}
html, body, [class*="css"], .stApp {{ font-family: 'Plus Jakarta Sans', sans-serif !important; background-color: var(--bg) !important; }}
.stApp p, .stApp span, .stApp label, .stApp li, .stApp h1, .stApp h2, .stApp h3, .stApp h4 {{ color: var(--text) !important; }}
.main-header {{ background: {_header_gr}; border-radius: 20px; padding: 36px 40px; margin-bottom: 28px; box-shadow: 0 8px 40px rgba(0,0,0,0.1); color: white; }}
.main-header h1, .main-header p {{ color: white !important; }}
[data-testid="stSidebar"] {{ background: {_sidebar_gr} !important; }}
[data-testid="stSidebar"] * {{ color: white !important; }}
.stButton > button {{ background: {_btn_gr} !important; color: white !important; border-radius: 12px !important; border: none !important; padding: 12px 28px !important; font-weight: 700 !important; }}
div[data-testid="metric-container"] {{ background: var(--card) !important; border-radius: 14px !important; padding: 18px !important; border: 1px solid var(--border) !important; }}
</style>
""", unsafe_allow_html=True)

# =========================
# 🧠 로직 파트 (유지)
# =========================
CACHE = {}
def _hash_key(prompt: str, model: str):
    return hashlib.md5(f"{model}:{prompt}".encode()).hexdigest()

def cached_call(call_fn, prompt: str, model: str, ttl=3600):
    key = _hash_key(prompt, model)
    if key in CACHE:
        data, ts = CACHE[key]
        if time.time() - ts < ttl: return data
    try:
        result = call_fn(prompt)
        CACHE[key] = (result, time.time())
        return result
    except: return ""

def call_gpt(prompt):
    if "openai_key" not in st.session_state: return ""
    client = OpenAI(api_key=st.session_state["openai_key"])
    def _call(p):
        res = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": p}],
            max_tokens=300
        )
        return res.choices[0].message.content.strip()
    return cached_call(_call, prompt, "gpt")

def call_gemini(prompt):
    if "gemini_key" not in st.session_state: return ""
    def _call(p):
        url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-pro:generateContent?key={st.session_state['gemini_key']}"
        body = {"contents": [{"parts": [{"text": p}]}]}
        res = requests.post(url, json=body)
        if res.status_code == 200: return res.json()["candidates"][0]["content"]["parts"][0]["text"]
        return ""
    return cached_call(_call, prompt, "gemini")

def build_variants(domain, brand):
    base = domain.split(".")[0]
    return list(set([base.lower(), domain.lower(), brand.lower(), brand.replace(" ", "").lower()]))

def is_brand_mentioned(text, variants):
    text = text.lower()
    for v in variants:
        if v in text: return True
        for word in text.split():
            if difflib.SequenceMatcher(None, word, v).ratio() > 0.75: return True
    return False

def generate_questions(brand, industry):
    prompt = f"{industry}에서 {brand} 관련 질문 5개 생성"
    res = call_gpt(prompt) or call_gemini(prompt)
    return [q.strip("- ").strip() for q in res.split("\n") if "?" in q][:5]

def run_simulation(question, domain, brand, n=20):
    variants = build_variants(domain, brand)
    use_gpt = "openai_key" in st.session_state
    use_gemini = "gemini_key" in st.session_state
    def single():
        prompt = f"질문: {question}\n답변:"
        hit, total = 0, 0
        if use_gpt:
            res = call_gpt(prompt)
            total += 1
            if is_brand_mentioned(res, variants): hit += 1
        if use_gemini:
            res = call_gemini(prompt)
            total += 1
            if is_brand_mentioned(res, variants): hit += 1
        return hit, total
    hits, total = 0, 0
    with concurrent.futures.ThreadPoolExecutor(max_workers=5) as ex:
        futures = [ex.submit(single) for _ in range(n)]
        for f in futures:
            try:
                h, t = f.result()
                hits += h
                total += t
            except: pass
    rate = round((hits / total * 100), 1) if total > 0 else 0
    return hits, total, rate

def run_strategy(question, brand):
    prompt = f"{brand}이 '{question}'에서 노출을 높이는 전략 3가지"
    return call_gpt(prompt) or call_gemini(prompt)

def save_word_report(data):
    doc = Document()
    doc.add_heading(f"AI Citation Report: {data['brand']}", 0)
    doc.add_paragraph(f"도메인: {data['domain']}")
    doc.add_paragraph(f"생성일시: {data['created_at']}")
    doc.add_paragraph("\n")
    
    for r in data["results"]:
        doc.add_heading(f"❓ {r['question']}", level=1)
        doc.add_paragraph(f"인용 횟수: {r['hits']}회")
        doc.add_paragraph(f"총 시도: {r['total']}회")
        doc.add_paragraph(f"점유율: {r['rate']}%")
        doc.add_paragraph("\n")
    
    # 메모리에서 바로 다운로드 가능
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# =========================
# 🎨 사이드바 UI
# =========================
with st.sidebar:
    st.markdown("""<div style='text-align:center; padding:20px 0;'><span style='font-size:2.5rem;'>🔍</span><h2 style='margin:0;'>AI Citation</h2></div>""", unsafe_allow_html=True)
    
    _mode_label = "☀️ 라이트 모드로 전환" if _dark else "🌙 다크 모드로 전환"
    if st.button(_mode_label, use_container_width=True):
        st.session_state["dark_mode"] = not st.session_state["dark_mode"]
        st.rerun()

    st.markdown("### 🔐 API 설정")
    openai_key = st.text_input("OpenAI API Key", type="password", placeholder="sk-...")
    gemini_key = st.text_input("Gemini API Key", type="password", placeholder="AIza...")
    if openai_key: st.session_state["openai_key"] = openai_key
    if gemini_key: st.session_state["gemini_key"] = gemini_key
    
    st.markdown("---")
    st.caption("AI Citation Analyzer v22.0")

# =========================
# 🚀 메인 화면 UI
# =========================
st.markdown(f"""
<div class="main-header">
    <h1>🔍 AI 검색 점유율 분석기</h1>
    <p>GPT & Gemini 엔진에서 내 브랜드가 얼마나 인용되는지 시뮬레이션합니다</p>
</div>
""", unsafe_allow_html=True)

tab1, tab2 = st.tabs(["🤖 분석 자동화", "📅 AI 인용 히스토리"])

with tab1:
    st.markdown("""<div class='result-card'><h4>🎯 분석 대상 설정</h4>""", unsafe_allow_html=True)
    col_in1, col_in2, col_in3 = st.columns(3)
    brand_val = col_in1.text_input("🏷️ 브랜드명", "에이바헤어")
    domain_val = col_in2.text_input("🌐 도메인", "avahair.co.kr")
    industry_val = col_in3.text_input("🏭 산업군", "미용실 프랜차이즈")
    st.markdown("</div>", unsafe_allow_html=True)

    if st.button("분석 시작", use_container_width=True):
        if "openai_key" not in st.session_state and "gemini_key" not in st.session_state:
            st.error("좌측 사이드바에서 API 키를 먼저 입력하세요.")
        else:
            report_data = {"brand": brand_val, "domain": domain_val, "created_at": str(datetime.now()), "results": []}
            questions = generate_questions(brand_val, industry_val)
            
            with st.spinner("AI 엔진 시뮬레이션 및 분석 중..."):
                for q in questions:
                    with st.container():
                        st.markdown(f"""<div class='metric-card'><h3 style='margin:0;'>❓ {q}</h3>""", unsafe_allow_html=True)
                        hits, total, rate = run_simulation(q, domain_val, brand_val)
                        
                        m1, m2, m3 = st.columns(3)
                        m1.metric("인용 횟수", f"{hits}회")
                        m2.metric("총 시도", f"{total}회")
                        m3.metric("점유율 (%)", f"{rate}%")
                        
                        st.markdown("**🧠 GEO 최적화 전략:**")
                        strategy = run_strategy(q, brand_val)
                        st.info(strategy)
                        
                        report_data["results"].append({"question": q, "hits": hits, "total": total, "rate": rate})
                        st.markdown("</div><br>", unsafe_allow_html=True)

            # Word(.docx) 다운로드
            word_buffer = save_word_report(report_data)
            st.success(f"✅ 분석 완료: Word 리포트 생성 완료")
            st.download_button(
                "📥 결과 Word 다운로드",
                word_buffer,
                file_name=f"report_{brand_val}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

with tab2:
    st.info("로그 파일을 업로드하여 AI 엔진별 브랜드 누적 인용 횟수를 시각화하는 기능입니다. (준비 중)")

st.markdown("<div style='height:40px'></div><div style='text-align:center; color:var(--text_muted); font-size:0.8rem; border-top:1px solid var(--border); padding-top:20px;'>🔍 AI Citation Analyzer | GPT & Gemini AI 검색 점유율 분석 도구</div>", unsafe_allow_html=True)
